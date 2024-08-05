import math
import ffmpeg
from faster_whisper import WhisperModel
from tqdm import tqdm
from docx import Document

from timecode import extract_time_code_and_framerate
from timecode import add_seconds_to_time_code

def format_time(seconds):

    hours = math.floor(seconds / 3600)
    seconds %= 3600
    minutes = math.floor(seconds / 60)
    seconds %= 60
    milliseconds = round((seconds - math.floor(seconds)) * 1000)
    seconds = math.floor(seconds)
    formatted_time = f"{hours:02d}:{minutes:02d}:{seconds:01d},{milliseconds:03d}"

    return formatted_time

def extract_audio(input_video, input_video_name, output_dir):
    extracted_audio = f"{output_dir}/audio-{input_video_name}.wav"
    ffmpeg.input(input_video).output(extracted_audio).run()

    return extracted_audio

def transcribe(audio: str, model: str, device: str, compute_type, language: str):
    model = WhisperModel(model, device=device, compute_type=compute_type)
    segments, info = model.transcribe(audio, language=language)
    language = info[0]
    print("Transcription language", language)
    segments = list(tqdm(segments, desc="Transcribing", unit="segment"))
    return segments


def generate_srt_file(segments, input_video_name, output_dir):

    subtitle_file = f"{output_dir}/sub-{input_video_name}.srt"
    text = ""
    for index, segment in enumerate(segments):
        segment_start = format_time(segment.start)
        segment_end = format_time(segment.end)
        text += f"{str(index+1)} \n"
        text += f"{segment_start} --> {segment_end} \n"
        text += f"{segment.text} \n"
        text += "\n"
        
    f = open(subtitle_file, "w")
    f.write(text)
    f.close()

    return subtitle_file

def generate_timecode_file(segments, input_video, input_video_name, output_dir):
    
    timecode_file = f"{output_dir}/timecode-{input_video_name}.txt"
    start_timecode, fps = extract_time_code_and_framerate(input_video)
    text = ""
    for index, segment in enumerate(segments):
               
            if index == 0 and segment.start < 5:
                timecode_start = start_timecode
            else:
                timecode_start = add_seconds_to_time_code(time_code=start_timecode, add_seconds=segment.start, frame_rate=fps)
            timecode_end = add_seconds_to_time_code(time_code=start_timecode, add_seconds=segment.end, frame_rate=fps)
            text += f"{str(index+1)}\n"
            text += f"{timecode_start} --> {timecode_end}\n"
            text += f"{segment.text}\n"
            text += "\n"
            
    
    f = open(timecode_file, "w")
    f.write(text)
    f.close()

    return timecode_file

def generate_doc_file(segments, input_video, input_video_name, output_dir):
    doc_file = f"{output_dir}/word-{input_video_name}.docx"
    
    doc = Document()
    doc.add_heading(f"{input_video_name}", 0)
    doc.add_heading(f"Automatically created Docx file from video speech.", level=1)
    par = doc.add_paragraph("The time of the video is ")
    par.add_run("bolded").bold = True
    par.add_run(" and timecode is in ")
    par.add_run("italics").italic = True
    par.add_run(".")

    start_timecode, fps = extract_time_code_and_framerate(input_video)
    text = doc.add_paragraph("")
    for index, segment in enumerate(segments):
               
            if index == 0 and segment.start < 5:
                timecode_start = start_timecode
            else:
                timecode_start = add_seconds_to_time_code(time_code=start_timecode, add_seconds=segment.start, frame_rate=fps)

            if index == 0 or index % 10 == 0:
                text.add_run("\n")
                text.add_run(f"{format_time(segment.start)}").bold = True
                text.add_run(" | ")
                text.add_run(f"{timecode_start}\n").italic = True
                text.add_run(f"{segment.text}\n")
            else:
                text.add_run(f"{segment.text}\n")
    doc.save(doc_file)
    return doc_file

def generate_doc_file_no_timecode(segments, input_video_name, output_dir):
    doc_file = f"{output_dir}/word-no-timecode-{input_video_name}.docx"
    
    doc = Document()
    doc.add_heading(f"{input_video_name}", 0)
    doc.add_heading(f"Automatically created Docx file from video speech.", level=1)
    par = doc.add_paragraph("The time of the video is ")
    par.add_run("bolded").bold = True
    par.add_run(", but the timecode wasn't available.")

    text = doc.add_paragraph("")
    for index, segment in enumerate(segments):

            if index == 0 or index % 10 == 0:
                text.add_run("\n")
                text.add_run(f"{format_time(segment.start)}\n").bold = True
                text.add_run(f"{segment.text}\n")
            else:
                text.add_run(f"{segment.text}\n")
    doc.save(doc_file)
    return doc_file
